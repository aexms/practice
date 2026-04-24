"""
Сборка PRACTICE_INTERNSHIP_REPORT.docx из PRACTICE_INTERNSHIP_REPORT.md (упрощённый разбор).

Установка: pip install -r practice/requirements-docx.txt
Запуск из каталога diplom: python practice/build_practice_report_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Установите: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def _strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def _add_runs_with_bold(paragraph, text: str) -> None:
    text = text.rstrip("\n")
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**") and len(p) > 4:
            paragraph.add_run(p[2:-2]).bold = True
        else:
            paragraph.add_run(_strip_md_bold(p))


def _parse_table_lines(lines: list[str]) -> list[list[str]] | None:
    if not lines or not lines[0].strip().startswith("|"):
        return None
    rows: list[list[str]] = []
    for line in lines:
        s = line.strip()
        if not s.startswith("|"):
            break
        if re.match(r"^\|?[\s\-:|]+\|?$", s.replace(" ", "")):
            continue
        cells = [c.strip() for c in s.split("|")]
        cells = [c for c in cells if c != ""]
        if cells:
            rows.append(cells)
    return rows if rows else None


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                block.append(lines[j])
                j += 1
            rows = _parse_table_lines(block)
            i = j
            if rows and len(rows) >= 1:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncol)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        text = row[ci] if ci < len(row) else ""
                        tbl.cell(ri, ci).text = _strip_md_bold(text)
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            nhash = len(m.group(1))
            text = m.group(2).strip()
            text = _strip_md_bold(text)
            # Word: level 0 ≈ заголовок документа; ## → Heading 1, ### → Heading 2 …
            word_level = max(0, min(nhash - 1, 3))
            doc.add_heading(text, level=word_level)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs_with_bold(p, stripped)
        i += 1

    p = doc.add_paragraph()
    r = p.add_run(
        "Примечание: при необходимости в Word задайте стили заголовков (Заголовок 1–3) "
        "и поля страницы по требованиям кафедры."
    )
    r.italic = True

    doc.save(docx_path)
    print("Written", docx_path)


def main() -> None:
    here = Path(__file__).resolve().parent
    md = here / "PRACTICE_INTERNSHIP_REPORT.md"
    out = here / "PRACTICE_INTERNSHIP_REPORT.docx"
    if not md.is_file():
        print("Не найден", md, file=sys.stderr)
        sys.exit(1)
    md_to_docx(md, out)


if __name__ == "__main__":
    main()
